"""ייצוא טקסט מקובץ PDF כמשימת רקע עם דיווח התקדמות.

שלושה מקורות טקסט:
- ``saved``: הטקסט השמור בקטלוג (כולל OCR שכבר בוצע) - מיידי, ללא סריקה.
- ``text``: חילוץ חכם של שכבת הטקסט (pdfplumber) - מהיר, בלי OCR.
- ``ocr``: סריקת OCR מחדש של העמודים במנוע נבחר - איטי, מדויק לסרוקים.

הייצוא רץ ב-thread יחיד; ההתקדמות נשאלת ב-polling (עמוד X מתוך Y) וניתנת
לביטול. התוצאה נכתבת לנתיב שבחר המשתמש (TXT/DOCX), או נשמרת בזיכרון
להורדת דפדפן כשאין דיאלוג שמירה נייטיבי.
"""

from __future__ import annotations

import threading
import time
from pathlib import Path
from typing import Dict, List, Optional

from .catalog import catalog
from .logging_setup import get_logger

log = get_logger("export")

_lock = threading.Lock()
_cancel = threading.Event()
_thread: Optional[threading.Thread] = None
_result_text: str = ""

status: Dict = {
    "running": False,
    "source": "",
    "path": "",
    "name": "",
    "page": 0,
    "pages": 0,
    "target": "",
    "error": "",
    "done": False,
    "result_available": False,
}


def _set(**kw) -> None:
    status.update(kw)


def write_text_file(target: Path, text: str, fmt: str = "txt") -> Path:
    """כתיבת טקסט לקובץ TXT או DOCX (עם כיווניות RTL)."""
    target.parent.mkdir(parents=True, exist_ok=True)
    if fmt == "docx" or target.suffix.lower() == ".docx":
        _write_docx(target, text)
        return target if target.suffix else target.with_suffix(".docx")
    if not target.suffix:
        target = target.with_suffix(".txt")
    target.write_text(text, encoding="utf-8-sig")
    return target


def _write_docx(target: Path, text: str) -> None:
    """שמירת טקסט כקובץ Word (docx) עם כיווניות RTL."""
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement

    doc = docx.Document()
    for line in text.splitlines() or [""]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        try:
            pPr = p._p.get_or_add_pPr()
            pPr.append(OxmlElement("w:bidi"))
        except Exception:
            pass
    if not target.suffix:
        target = target.with_suffix(".docx")
    doc.save(str(target))


def is_running() -> bool:
    return _thread is not None and _thread.is_alive()


def get_status() -> Dict:
    return dict(status)


def get_result() -> str:
    """הטקסט האחרון שיוצא ללא נתיב יעד (להורדת דפדפן)."""
    return _result_text


def cancel() -> None:
    _cancel.set()


def start_export(
    path: str,
    target: str = "",
    source: str = "saved",
    engine_id: str = "",
    fmt: str = "txt",
    page_from: int = 0,
    page_to: int = 0,
) -> Dict:
    """מפעיל ייצוא ברקע. מחזיר {"started": True} או {"error": ...}."""
    global _thread
    with _lock:
        if is_running():
            return {"error": "ייצוא אחר כבר רץ - המתינו לסיומו"}
        p = Path(path)
        if not p.exists():
            return {"error": "הקובץ לא נמצא"}
        _cancel.clear()
        _set(
            running=True, source=source, path=path, name=p.stem,
            page=0, pages=0, target=target, error="", done=False,
            result_available=False,
        )
        _thread = threading.Thread(
            target=_run_export,
            args=(p, target, source, engine_id, fmt, page_from, page_to),
            daemon=True,
        )
        _thread.start()
        return {"started": True}


def _run_export(
    p: Path, target: str, source: str, engine_id: str,
    fmt: str, page_from: int, page_to: int,
) -> None:
    global _result_text
    try:
        if source == "saved":
            text = _from_catalog(p, page_from, page_to)
        elif source == "ocr":
            text = _from_ocr(p, engine_id, page_from, page_to)
        else:
            text = _from_text_layer(p, page_from, page_to)

        if _cancel.is_set():
            _set(running=False, error="הייצוא בוטל", done=False)
            return
        if not text.strip():
            _set(running=False, error="לא נמצא טקסט לייצוא", done=False)
            return

        if target:
            written = write_text_file(Path(target), text, fmt)
            log.info("ייצוא טקסט הושלם: %s -> %s (%d תווים)", p.name, written, len(text))
            _set(running=False, done=True, target=str(written))
        else:
            # אין דיאלוג שמירה נייטיבי - התוצאה תרד דרך הדפדפן
            _result_text = text
            _set(running=False, done=True, result_available=True)
    except Exception as exc:
        log.warning("ייצוא טקסט נכשל עבור %s: %s", p, exc)
        _set(running=False, error=f"הייצוא נכשל: {exc}", done=False)


def _page_range(total: int, page_from: int, page_to: int) -> List[int]:
    """רשימת עמודים 1-based בטווח המבוקש (0 = ללא הגבלה)."""
    lo = max(1, page_from or 1)
    hi = min(total, page_to or total)
    return list(range(lo, hi + 1))


def _from_catalog(p: Path, page_from: int, page_to: int) -> str:
    """הטקסט השמור בקטלוג (מהאינדוקס/OCR שכבר בוצע), לפי טווח עמודים."""
    row = catalog.get_file_by_path(str(p))
    if row is None or not row["full_text"]:
        raise RuntimeError(
            "אין טקסט שמור לקובץ זה - הוסיפו את התיקייה לאינדוקס והמתינו לסיום, "
            "או בחרו מקור אחר"
        )
    full_text = row["full_text"]
    segs = catalog.get_segments(row["id"])
    if not segs or (not page_from and not page_to):
        _set(page=row["page_count"] or 0, pages=row["page_count"] or 0)
        return full_text
    wanted = set(_page_range(row["page_count"] or len(segs), page_from, page_to))
    parts = [
        full_text[s["char_start"]:s["char_end"]]
        for s in segs
        if s["page"] in wanted
    ]
    if not parts:
        raise RuntimeError("אין טקסט שמור בטווח העמודים שנבחר")
    _set(page=len(parts), pages=len(parts))
    return "\n\n".join(parts)


def _from_text_layer(p: Path, page_from: int, page_to: int) -> str:
    """חילוץ חכם של שכבת הטקסט, עמוד-עמוד עם התקדמות."""
    import pdfplumber

    from .extractors import pdf_smart

    out: List[str] = []
    with pdfplumber.open(str(p)) as pdf:
        pages = _page_range(len(pdf.pages), page_from, page_to)
        _set(pages=len(pages))
        for i, num in enumerate(pages, start=1):
            if _cancel.is_set():
                return ""
            try:
                text = pdf_smart.extract_page_text(pdf.pages[num - 1])
            except Exception as exc:
                log.debug("חילוץ עמוד %d נכשל: %s", num, exc)
                text = ""
            if text:
                out.append(text)
            _set(page=i)
    return "\n\n".join(out)


def _from_ocr(p: Path, engine_id: str, page_from: int, page_to: int) -> str:
    """סריקת OCR של העמודים בטווח, במנוע הייצוא הנבחר."""
    import pypdfium2 as pdfium

    from . import hebrew_bidi
    from .extractors import ocr_engines

    engine = ocr_engines.get_export_engine(engine_id)
    if not engine.available():
        raise RuntimeError(f"מנוע ה-OCR '{engine_id or engine.id}' אינו זמין")

    pdf = pdfium.PdfDocument(str(p))
    try:
        pages = _page_range(len(pdf), page_from, page_to)
        _set(pages=len(pages))
        scale = engine.render_dpi() / 72.0
        out: List[str] = []
        done = 0
        step = max(1, engine.pdf_batch)
        for start in range(0, len(pages), step):
            if _cancel.is_set():
                return ""
            batch = pages[start : start + step]
            images = [pdf[num - 1].render(scale=scale).to_pil() for num in batch]
            texts = engine.ocr_images(images)
            for txt in texts:
                fixed, _ = hebrew_bidi.fix_visual_order(txt)
                out.append(fixed)
            done += len(batch)
            _set(page=done)
        return "\n\n".join(t for t in out if t.strip())
    finally:
        pdf.close()
        # שחרור משאבי המנוע (למשל worker של Surya) - רק אם תור ה-OCR הרקעי
        # ריק, כדי לא להפיל worker שהאינדוקס עדיין משתמש בו
        try:
            if catalog.count_pending_ocr() == 0:
                engine.idle()
        except Exception:
            pass
