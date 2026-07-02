"""חילוץ מסמכי Word (.docx) כולל פסקאות וטבלאות."""

from __future__ import annotations

from pathlib import Path

from ..logging_setup import get_logger
from .base import ExtractResult, Page

log = get_logger("extract.docx")


def extract(path: Path) -> ExtractResult:
    import docx  # python-docx

    document = docx.Document(str(path))
    parts: list[str] = []

    for para in document.paragraphs:
        if para.text and para.text.strip():
            parts.append(para.text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))

    text = "\n".join(parts)
    log.debug("חולצו %d פסקאות/שורות מ-%s", len(parts), path.name)
    return ExtractResult(pages=[Page(number=1, text=text)], source="extracted")
